# -*- coding: utf-8 -*-
import docx
from docx.shared import Cm  # 加入可調整的 word 單位
import os
import itertools

doc = docx.Document("test.docx")

# 獲取所有圖片
# 取得當前目錄的所有檔案名稱(注意：結果會包含程式本身)
filenames = os.listdir(".")

def is_imag(filename):
    for ext in [".png", ".jpg"]:
        if filename.endswith(ext):
            return True
    return False

images = filter(is_imag, filenames)
# 建立一個新的圖片清單
newlist = list(images)
# print(newlist)

## 插入全部圖片
# for image in newlist:
#     # 新增資料夾內的照片
#     doc.add_picture(image, width=Cm(18))

# textlist=[]
# for i, p in enumerate(doc.paragraphs): # 遍历所有的段落
#     # print(str(i) + ":"+ str(p.text))
#     # 所有的標題
#     style_name=p.style.name
#     if(p.text!="" and p.text!="\n" and p.text!="\t" and style_name=='Normal'):
#         textlist.append(p.text)
# print(textlist)

textlist=[]
for p in doc.paragraphs:
    style_name=p.style.name
    # print(p.text)
    if(p.text!="" and p.text!="\n" and p.text!="\t" and style_name=='Normal'):
        textlist.append(p.text)
# print(textlist)



# 指定位置插入圖片
for i, p in enumerate(doc.paragraphs): # 遍历所有的段落
    if p.text in textlist:
        print(p.text)
        print(newlist[i-4])
        p.runs[-1].add_break()  # 添加一个折行
        p.runs[-1].add_picture(newlist[i-4], width=Cm(18)) # 在runs的最后一段文字后添加图片
        p.runs[-1].add_break(docx.enum.text.WD_BREAK.PAGE)

# for i, p in itertools.zip_longest(newlist,doc.paragraphs):
#     if (p.text in textlist):
#         print(p.text)
#         for i in newlist:
#             print(i)
#     #     print(p.text)
#     #     print(i)
#         # p.runs[-1].add_picture(i, width=Cm(18)) # 在runs的最后一段文字后添加图片



# for i, p in enumerate(doc.paragraphs): # 遍历所有的段落
#     if p.text in textlist:
#         # print("進入了")
#         # p.runs[-1].add_picture(newlist[i], width=Cm(18)) # 在runs的最后一段文字后添加图片

#         # p.runs[-1].add_break()  # 添加一个折行
#         # p.runs[-1].add_break()  # 添加一个折行
#         p.runs[-1].add_picture("Screenshots2022-01-26_15_26_07舊客戶資料-證件號碼x.png", width=Cm(18)) # 在runs的最后一段文字后添加图片

# # 指定位置插入圖片
# for i, p in enumerate(doc.paragraphs): # 遍历所有的段落
#     print(str(i) + ":"+ str(p.text))
#     if len(p.text) != 0:
#         for i in range(len(p.runs)): # p.runs代表p这个段落下所有文字的列表
#             print(p.runs[i].text)  # 当打印时，发现p.runs把段落自动分解了
#     print("---------------------------------------------------------------")
#     if "1.保單號碼：無須對位" in p.text:
#         print("進入了")
#         p.runs[-1].add_break()  # 添加一个折行
#         # p.runs[-1].add_break()  # 添加一个折行
#         p.runs[-1].add_picture("Screenshots2022-01-26_15_26_07舊客戶資料-證件號碼x.png", width=Cm(18)) # 在runs的最后一段文字后添加图片


doc.save('AI 技術可以讓隱藏於暗處的物品現形.docx')
